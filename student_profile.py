from docx import Document
from docx.shared import Pt

# Create a new Word Document
doc = Document()
doc.add_heading('STUDENT PROFILE', 0)

# Set a clean, professional font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Define the fields
fields = [
    ("1. Student Name (As per Aadhaar)", "GOWSHIK J"),
    ("2. Class", "VI - A"),
    ("3. Admission No (SID)", "6071242"),
    ("4. Student Name in Tamil", "கௌசிக் ஜ"),
    ("5. Student Aadhaar", "3464 3282 4641"),
    ("6. Date of Birth", "09-04-2014"),
    ("7. Gender", "Male"),
    ("8. Blood Group", "O+ve"),
    ("9. Mother Tongue", "Tamil"),
    ("10. Disability Group", "No"),
    ("11. Religion", "Hindu"),
    ("12. Community", "BC (Backward Class)"),
    ("13. Caste", "Kongu Vellalar"),
    ("14. Father’s Name in English", "JANARTHANAN PONNUSAMY"),
    ("15. Father’s Name in Tamil", "ஜனார்த்தனன் பொன்னுச்சாமி"),
    ("16. Father's Education Level", "B.Sc. Computer Science"),
    ("17. Father's Occupation", "Business"),
    ("18. Mother's Education Level", "Diploma"),
    ("19. Mother's Occupation", "Business"),
    ("20. Guardian's Name in English", "P. KAMALAM"),
    ("21. Guardian's Name in Tamil", "P. கமலம்"),
    ("22. Guardian's Education Level", "Primary Level"),
    ("23. Parents' Annual Income", "4 Lakhs"),
    ("24. Mobile Number", "98426 83539"),
    ("25. Email ID", "janabsc@yahoo.com"),
    ("26. Door No / Building No", "9"),
    ("27. Street Name / Area Name", "Kasiyanang Kattu Thottam,\n    Kallumadai"),
    ("28. City / Village Name", "Kolappalur (Po),\n    Gobichettipalayam (Tk)"),
    ("29. District", "Erode"),
    ("30. Pincode", "638456"),
    ("31. Class Section", "VI - A"),
    ("32. Date of Joining", "(Not provided)"),
    ("33. Medium of Instruction", "English"),
    ("34. Admission No", "6071242"),
    ("35. Mother's Name (As per Aadhaar)", "J. Yamunadevi"),
    ("36. Mother's Phone Number", "84382 96668"),
    ("37. Previous School Name", "Spring Mount Valley School,\n    Perundurai"),
    ("38. EMIS Number", "(Not provided)"),
    ("39. TC Submitted", "No")
]

# Add each line to the document
for field, value in fields:
    para = doc.add_paragraph()
    para.add_run(f"{field:45}: ").bold = True
    para.add_run(value)

# Save the file
doc.save("Student_Profile.docx")
